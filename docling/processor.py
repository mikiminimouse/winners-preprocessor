"""
Модуль для обработки документов через Docling и другие библиотеки.
Обрабатывает: PDF (OCR и text extraction), DOCX, HTML, Layout Analysis, Table Extraction.
"""
import os
import time
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime
import hashlib

try:
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions, VlmPipelineOptions
    try:
        # Попытка импортировать VLM компоненты (может быть недоступно в старых версиях)
        from docling.datamodel.pipeline_options import ApiVlmOptions
        from docling.pipeline.vlm_pipeline import VlmPipeline
        VLM_AVAILABLE = True
    except ImportError:
        VLM_AVAILABLE = False
        logger = logging.getLogger(__name__)
        logger.info("VLM pipeline not available, using standard OCR")
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    VLM_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("docling-core not available, will use fallback libraries")

# Опциональные импорты для fallback обработки (не нужны для удаленного VLM)
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

# Cloud.ru S3 клиент (опционально)
try:
    # Импортируем из той же папки
    import importlib.util
    s3_client_path = Path(__file__).parent / "s3_client.py"
    if s3_client_path.exists():
        spec = importlib.util.spec_from_file_location("s3_client", s3_client_path)
        s3_client_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(s3_client_module)
        CloudRuS3Client = s3_client_module.CloudRuS3Client
        S3_AVAILABLE = True
    else:
        CloudRuS3Client = None
        S3_AVAILABLE = False
except ImportError as e:
    CloudRuS3Client = None
    S3_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning(f"Cloud.ru S3 client not available: {e}")
except Exception as e:
    CloudRuS3Client = None
    S3_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning(f"Cloud.ru S3 client not available: {e}")

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

import io

logger = logging.getLogger(__name__)

# MongoDB для сохранения метрик
MONGO_SERVER = os.environ.get("MONGO_SERVER", "mongodb:27017")
MONGO_USER = os.environ.get("MONGO_METADATA_USER", "docling_user")
MONGO_PASSWORD = os.environ.get("MONGO_METADATA_PASSWORD", "password")
MONGO_DB = os.environ.get("MONGO_METADATA_DB", "docling_metadata")

_mongo_client = None


def get_mongo_client():
    """Получает MongoDB клиент."""
    global _mongo_client
    if _mongo_client is None:
        try:
            from pymongo import MongoClient
            url = f'mongodb://{MONGO_USER}:{MONGO_PASSWORD}@{MONGO_SERVER}/?authSource=admin'
            _mongo_client = MongoClient(url)
        except Exception as e:
            logger.warning(f"Could not connect to MongoDB: {e}")
            return None
    return _mongo_client


def save_processing_metrics(metrics: Dict[str, Any]) -> bool:
    """Сохраняет метрики обработки в MongoDB."""
    client = get_mongo_client()
    if not client:
        return False
    
    try:
        db = client[MONGO_DB]
        collection = db["processing_metrics"]
        collection.insert_one(metrics)
        return True
    except Exception as e:
        logger.error(f"Error saving metrics: {e}")
        return False


def save_processing_error(error: Dict[str, Any]) -> bool:
    """Сохраняет ошибку обработки в MongoDB."""
    client = get_mongo_client()
    if not client:
        return False
    
    try:
        db = client[MONGO_DB]
        collection = db["processing_errors"]
        collection.insert_one(error)
        return True
    except Exception as e:
        logger.error(f"Error saving error: {e}")
        return False


def get_cached_result(file_hash: str) -> Optional[Dict[str, Any]]:
    """Получает кэшированный результат обработки по SHA256 хешу файла."""
    client = get_mongo_client()
    if not client:
        return None
    
    try:
        db = client[MONGO_DB]
        collection = db["processing_cache"]
        
        # Ищем кэш по хешу файла
        cached = collection.find_one({"file_hash": file_hash})
        if cached:
            cached.pop("_id", None)
            logger.info(f"Cache hit for file hash: {file_hash[:16]}...")
            return cached
        return None
    except Exception as e:
        logger.warning(f"Error getting cache: {e}")
        return None


def save_to_cache(file_hash: str, result: Dict[str, Any], ttl_days: int = 30) -> bool:
    """Сохраняет результат обработки в кэш с TTL."""
    client = get_mongo_client()
    if not client:
        return False
    
    try:
        from datetime import timedelta
        
        db = client[MONGO_DB]
        collection = db["processing_cache"]
        
        # Добавляем TTL - результаты будут автоматически удалены через ttl_days дней
        expires_at = datetime.utcnow() + timedelta(days=ttl_days)
        
        cache_entry = {
            "file_hash": file_hash,
            "result": result,
            "created_at": datetime.utcnow().isoformat(),
            "expires_at": expires_at.isoformat()
        }
        
        # Используем upsert для обновления существующего кэша
        collection.update_one(
            {"file_hash": file_hash},
            {"$set": cache_entry},
            upsert=True
        )
        
        # Создаем TTL индекс если не существует
        try:
            collection.create_index("expires_at", expireAfterSeconds=0)
        except Exception:
            pass  # Индекс уже существует
        
        logger.info(f"Result cached for file hash: {file_hash[:16]}... (expires in {ttl_days} days)")
        return True
    except Exception as e:
        logger.warning(f"Error saving to cache: {e}")
        return False


def calculate_file_hash(file_path: Path) -> str:
    """Вычисляет SHA256 хеш файла."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


# Глобальные переменные для кэширования конвертера
_cloudru_converter = None
_cloudru_converter_config = None

class DocumentProcessor:
    """Класс для обработки документов различных типов."""
    
    def __init__(self, unit_id: str, file_info: Dict[str, Any], output_dir: Path):
        self.unit_id = unit_id
        self.file_info = file_info
        self.file_path = Path(file_info["path"])
        self.output_dir = output_dir
        self.route = file_info.get("route", "unknown")
        self.detected_type = file_info.get("detected_type", "unknown")
        
        # Определение размера файла для оптимизации
        file_size = self.file_path.stat().st_size if self.file_path.exists() else 0
        self.is_large_file = file_size > 50 * 1024 * 1024  # 50 MB
        
        # Метрики обработки
        self.metrics = {
            "unit_id": unit_id,
            "file_id": file_info.get("file_id", ""),
            "file_name": file_info.get("original_name", ""),
            "route": self.route,
            "detected_type": self.detected_type,
            "processing_times": {},
            "file_stats": {
                "input_size": file_size,
                "is_large_file": self.is_large_file,
                "pages": 0,
                "tables_extracted": 0,
                "images_extracted": 0
            },
            "status": "processing",
            "error": None,
            "cached": False,
            "created_at": datetime.utcnow().isoformat()
        }
        
        # Проверка использования Cloud.ru VLM
        self.use_cloudru_vlm = os.getenv("USE_CLOUDRU_VLM", "false").lower() == "true"
        if self.use_cloudru_vlm:
            self.cloudru_api_url = os.getenv("CLOUDRU_API_URL")
            self.cloudru_api_key = os.getenv("CLOUDRU_API_KEY")
            self.cloudru_model = os.getenv("CLOUDRU_MODEL", "model-run-l4gj9-duty")
            logger.info(f"Using Cloud.ru VLM: {self.cloudru_model}")
        
        # Инициализация Cloud.ru S3 клиента (если доступен)
        self.use_s3 = os.getenv("USE_CLOUDRU_S3", "false").lower() == "true"
        self.s3_client = None
        if self.use_s3 and S3_AVAILABLE and CloudRuS3Client:
            try:
                # Cloud.ru требует формат tenant_id:key_id для access_key_id
                # Если CLOUDRU_S3_ACCESS_KEY уже содержит формат tenant_id:key_id, используем его
                # Иначе формируем из отдельных переменных
                access_key = os.getenv("CLOUDRU_S3_ACCESS_KEY")
                if not access_key or ":" not in access_key:
                    tenant_id = os.getenv("CLOUDRU_S3_TENANT_ID")
                    key_id = os.getenv("CLOUDRU_S3_KEY_ID")
                    if tenant_id and key_id:
                        access_key = f"{tenant_id}:{key_id}"
                
                self.s3_client = CloudRuS3Client(
                    endpoint_url=os.getenv("CLOUDRU_S3_ENDPOINT", "https://s3.cloud.ru"),
                    access_key_id=access_key,
                    secret_access_key=os.getenv("CLOUDRU_S3_SECRET_KEY"),
                    bucket_name=os.getenv("CLOUDRU_S3_BUCKET", "bucket-winners223"),
                    region_name=os.getenv("CLOUDRU_S3_REGION", "ru-central-1")
                )
                logger.info("Cloud.ru S3 client initialized successfully")
            except Exception as e:
                logger.warning(f"Could not initialize Cloud.ru S3 client: {e}")
                self.use_s3 = False
        
        # Инициализация Docling converter если доступен
        if DOCLING_AVAILABLE:
            try:
                # Проверяем использование Cloud.ru VLM
                if self.use_cloudru_vlm and self.cloudru_api_url and self.cloudru_api_key:
                    self.converter = self._init_cloudru_vlm_converter()
                    logger.info("Docling DocumentConverter with Cloud.ru VLM initialized successfully")
                # Проверяем доступность VLM адаптера (существующий код)
                elif VLM_AVAILABLE and os.getenv("USE_GRANITE_VLM", "false").lower() == "true":
                    # Конфигурация VLM Pipeline с удаленным Granite через адаптер
                    granite_adapter_url = os.getenv(
                        "GRANITE_ADAPTER_URL",
                        "http://granite_adapter:8000/process"
                    )
                    
                    logger.info(f"Initializing Docling with Granite VLM adapter: {granite_adapter_url}")
                    
                    api_opts = ApiVlmOptions(
                        url=granite_adapter_url,
                        method="POST",
                        headers={"Content-Type": "application/json"},
                        timeout=120
                    )
                    
                    vlm_opts = VlmPipelineOptions(
                        enable_remote_services=True,
                        api_vlm_options=api_opts
                    )
                    
                    pdf_opts = PdfFormatOption(
                        pipeline_cls=VlmPipeline,
                        pipeline_options=vlm_opts,
                        do_ocr=True
                    )
                    
                    self.converter = DocumentConverter(
                        format_options={
                            InputFormat.PDF: pdf_opts
                        }
                    )
                    logger.info("Docling DocumentConverter with Granite VLM initialized successfully")
                else:
                    # Стандартный converter без VLM
                    self.converter = DocumentConverter()
                    logger.info("Docling DocumentConverter initialized successfully (standard mode)")
                    
            except Exception as e:
                logger.warning(f"Could not initialize Docling converter: {e}")
                self.converter = None
        else:
            self.converter = None
    
    def _init_cloudru_vlm_converter(self):
        """Инициализация Docling конвертера с Cloud.ru VLM"""
        global _cloudru_converter, _cloudru_converter_config
        
        # Проверяем, нужно ли переинициализировать конвертер
        current_config = {
            "api_url": self.cloudru_api_url,
            "api_key": self.cloudru_api_key,
            "model": self.cloudru_model
        }
        
        if _cloudru_converter and _cloudru_converter_config == current_config:
            return _cloudru_converter
        
        try:
            from docling.datamodel.pipeline_options import VlmPipelineOptions
            from docling.datamodel.pipeline_options_vlm_model import ApiVlmOptions, ResponseFormat
            from docling.pipeline.vlm_pipeline import VlmPipeline
            
            # Конфигурация ApiVlmOptions для Cloud.ru согласно документации
            # https://docling-project.github.io/docling/examples/vlm_pipeline_api_model/
            # Используем правильный формат для OpenAI-compatible API
            vlm_options = ApiVlmOptions(
                url=self.cloudru_api_url,  # Endpoint Cloud.ru VLM API
                params={
                    "model": self.cloudru_model,
                    "max_tokens": 4096,  # Используем max_tokens для OpenAI-compatible API
                    "temperature": 0.0,  # Детерминированный вывод
                },
                headers={
                    "Authorization": f"Bearer {self.cloudru_api_key}",
                    "Content-Type": "application/json"
                },
                prompt="Extract all text from this document page. Preserve structure, tables, and formatting. Return in Russian.",
                timeout=180,  # Таймаут для больших документов
                scale=1.0,  # Масштаб изображения (1.0 для меньшего размера)
                response_format=ResponseFormat.MARKDOWN,  # Формат ответа
            )
            
            # Конфигурация VLM Pipeline (ТОЛЬКО удаленная обработка)
            pipeline_options = VlmPipelineOptions(
                enable_remote_services=True,
                vlm_options=vlm_options
            )
            
            # Создание конвертера (ТОЛЬКО с VLM pipeline, без локальных обработчиков)
            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(
                        pipeline_cls=VlmPipeline,
                        pipeline_options=pipeline_options,
                        do_ocr=False,  # Отключаем локальный OCR
                        do_layout=False,  # Отключаем локальный layout анализ
                        do_table_structure=False  # Отключаем локальную обработку таблиц
                    )
                }
            )
            
            # Кэшируем конвертер
            _cloudru_converter = converter
            _cloudru_converter_config = current_config
            
            return converter
            
        except Exception as e:
            logger.error(f"Failed to initialize Cloud.ru VLM converter: {e}")
            raise
    
    def process(self) -> Dict[str, Any]:
        """Обрабатывает документ согласно route."""
        start_time = time.time()
        
        try:
            # Проверка кэша перед обработкой
            file_hash = calculate_file_hash(self.file_path)
            cached_result = get_cached_result(file_hash)
            
            if cached_result:
                logger.info(f"Using cached result for {self.file_info.get('original_name', 'unknown')}")
                self.metrics["cached"] = True
                self.metrics["processing_times"]["total"] = time.time() - start_time
                self.metrics["status"] = "completed"
                self.metrics["completed_at"] = datetime.utcnow().isoformat()
                
                # Получаем output файлы из кэша
                cached_output = cached_result.get("result", {}).get("output_files", [])
                if cached_output:
                    self.metrics["output_files"] = cached_output
                    # Метрики из кэша
                    self.metrics.update(cached_result.get("result", {}).get("metrics", {}))
                    self.metrics["cached"] = True
                    
                    return {
                        "success": True,
                        "output_files": cached_output,
                        "metrics": self.metrics,
                        "cached": True
                    }
            
            logger.info(f"Processing {self.detected_type} file with route {self.route}")
            
            # Используем Cloud.ru VLM если настроено
            if self.use_cloudru_vlm:
                # Если доступен S3, используем его для передачи изображений
                if self.use_s3 and self.s3_client:
                    result = self._process_with_cloudru_vlm_s3()
                else:
                    result = self._process_with_cloudru_vlm()
            elif self.route == "pdf_scan":
                result = self._process_pdf_ocr()
            elif self.route == "pdf_text":
                result = self._process_pdf_text()
            elif self.route == "docx":
                result = self._process_docx()
            elif self.route == "html_text":
                result = self._process_html()
            else:
                logger.warning(f"Unknown route: {self.route}, using fallback")
                result = self._process_fallback()
            
            # Применяем Layout Analysis если не применен
            if not result.get("layout"):
                result["layout"] = self._extract_layout(result)
            
            # Сохранение результатов
            output_files = self._save_results(result)
            
            self.metrics["processing_times"]["total"] = time.time() - start_time
            
            # Добавляем время layout analysis если было применено отдельно
            if "layout_analysis" not in self.metrics["processing_times"]:
                self.metrics["processing_times"]["layout_analysis"] = 0
            
            self.metrics["status"] = "completed"
            self.metrics["completed_at"] = datetime.utcnow().isoformat()
            self.metrics["output_files"] = output_files
            
            # Добавляем processing_method в метрики
            self.metrics["processing_method"] = result.get("method", "unknown")
            
            # Сохранение метрик
            save_processing_metrics(self.metrics)
            
            # Сохранение в кэш
            result_data = {
                "output_files": output_files,
                "metrics": self.metrics,
                "result": {
                    "text": result.get("text", ""),
                    "tables": result.get("tables", []),
                    "metadata": result.get("metadata", {}),
                    "layout": result.get("layout", {})
                }
            }
            save_to_cache(file_hash, result_data)
            
            return {
                "success": True,
                "output_files": output_files,
                "metrics": self.metrics,
                "cached": False
            }
            
        except Exception as e:
            error_time = time.time() - start_time
            self.metrics["processing_times"]["total"] = error_time
            self.metrics["status"] = "failed"
            self.metrics["error"] = str(e)
            self.metrics["completed_at"] = datetime.utcnow().isoformat()
            
            # Сохранение ошибки
            error_data = {
                "unit_id": self.unit_id,
                "file_id": self.file_info.get("file_id", ""),
                "file_name": self.file_info.get("original_name", ""),
                "route": self.route,
                "error_type": type(e).__name__,
                "error_message": str(e),
                "stack_trace": self._get_stack_trace(e),
                "processing_time_before_error": error_time,
                "created_at": datetime.utcnow().isoformat()
            }
            save_processing_error(error_data)
            logger.error(f"Error processing document: {e}", exc_info=True)
            
            return {
                "success": False,
                "error": str(e),
                "metrics": self.metrics
            }
    
    def _process_pdf_ocr(self) -> Dict[str, Any]:
        """Обрабатывает PDF с OCR."""
        start_time = time.time()
        
        if DOCLING_AVAILABLE and self.converter:
            # Используем Docling для OCR
            result = self.converter.convert(str(self.file_path))
            doc_json = result.document.export_to_dict()
            
            # Извлечение текста и структуры
            text_content = []
            tables = []
            
            for item in doc_json.get("content", []):
                if item.get("type") == "text":
                    text_content.append(item.get("text", ""))
                elif item.get("type") == "table":
                    tables.append(item)
            
            self.metrics["processing_times"]["ocr"] = time.time() - start_time
            self.metrics["file_stats"]["pages"] = len(doc_json.get("pages", []))
            self.metrics["file_stats"]["tables_extracted"] = len(tables)
            
            return {
                "text": "\n".join(text_content),
                "tables": tables,
                "metadata": doc_json.get("metadata", {}),
                "layout": doc_json.get("layout", {}),
                "method": "docling_ocr"
            }
        else:
            # Fallback: используем pytesseract и pdfplumber
            return self._process_pdf_ocr_fallback()
    
    def _process_pdf_ocr_fallback(self) -> Dict[str, Any]:
        """Fallback OCR обработка PDF."""
        start_time = time.time()
        
        text_content = []
        tables = []
        pages_count = 0
        
        # Для больших файлов уменьшаем DPI и обрабатываем батчами
        dpi = 200 if self.is_large_file else 300
        batch_size = 5 if self.is_large_file else 10
        
        # Пробуем использовать pdf2image для OCR
        try:
            from pdf2image import convert_from_path
            
            # Для больших файлов обрабатываем по батчам
            if self.is_large_file:
                logger.info(f"Processing large PDF with OCR - will process in batches of {batch_size} pages")
                # Определяем общее количество страниц сначала
                if pdfplumber is None:
                    raise ImportError("pdfplumber not available")
                with pdfplumber.open(str(self.file_path)) as pdf:
                    pages_count = len(pdf.pages)
                
                # Обрабатываем по батчам
                for batch_start in range(0, pages_count, batch_size):
                    batch_end = min(batch_start + batch_size, pages_count)
                    logger.info(f"Processing OCR batch: pages {batch_start+1}-{batch_end}/{pages_count}")
                    
                    try:
                        images = convert_from_path(
                            str(self.file_path),
                            dpi=dpi,
                            first_page=batch_start + 1,
                            last_page=batch_end
                        )
                        
                        for idx, image in enumerate(images):
                            page_num = batch_start + idx + 1
                            try:
                                page_text = pytesseract.image_to_string(image, lang='rus+eng')
                                text_content.append(f"--- Page {page_num} ---\n{page_text}")
                            except Exception as e:
                                logger.warning(f"OCR failed for page {page_num}: {e}")
                                text_content.append(f"--- Page {page_num} ---\n(OCR failed)\n")
                    except Exception as e:
                        logger.error(f"Error processing OCR batch {batch_start+1}-{batch_end}: {e}")
            else:
                # Для небольших файлов - обычная обработка
                images = convert_from_path(str(self.file_path), dpi=dpi)
                pages_count = len(images)
                
                for i, image in enumerate(images):
                    try:
                        page_text = pytesseract.image_to_string(image, lang='rus+eng')
                        text_content.append(f"--- Page {i+1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"OCR failed for page {i+1}: {e}")
                        text_content.append(f"--- Page {i+1} ---\n(OCR failed)\n")
            
            # Извлечение таблиц через pdfplumber (только для первых страниц для больших файлов)
            try:
                if pdfplumber is None:
                    return {
                        "text": "\n".join(text_content),
                        "tables": [],
                        "metadata": {"pages": pages_count, "large_file_optimized": self.is_large_file},
                        "method": "pytesseract_fallback_no_pdfplumber"
                    }
                with pdfplumber.open(str(self.file_path)) as pdf:
                    pages_to_process = 50 if self.is_large_file else len(pdf.pages)
                    for i, page in enumerate(pdf.pages[:pages_to_process]):
                        page_tables = page.extract_tables()
                        if page_tables:
                            tables.extend(page_tables)
            except Exception as e:
                logger.warning(f"Table extraction failed: {e}")
            
            self.metrics["processing_times"]["ocr"] = time.time() - start_time
            self.metrics["file_stats"]["pages"] = pages_count
            self.metrics["file_stats"]["tables_extracted"] = len(tables)
            
            return {
                "text": "\n".join(text_content),
                "tables": tables,
                "metadata": {"pages": pages_count, "large_file_optimized": self.is_large_file},
                "method": "pytesseract_fallback"
            }
        except ImportError:
            # Если pdf2image недоступен, используем только pdfplumber
            if pdfplumber is None:
                logger.warning("pdf2image and pdfplumber not available, cannot process PDF")
                return {
                    "text": "",
                    "tables": [],
                    "metadata": {},
                    "method": "fallback_no_libraries"
                }
            logger.info("pdf2image not available, using pdfplumber only")
            with pdfplumber.open(str(self.file_path)) as pdf:
                pages_count = len(pdf.pages)
                
                # Для больших файлов обрабатываем постранично с логированием
                if self.is_large_file:
                    logger.info(f"Processing large PDF ({pages_count} pages) without OCR")
                
                for i, page in enumerate(pdf.pages):
                    if self.is_large_file and i % 10 == 0:
                        logger.info(f"Processing page {i+1}/{pages_count}")
                    
                    text = page.extract_text()
                    if text:
                        text_content.append(f"--- Page {i+1} ---\n{text}")
                    
                    # Таблицы только для первых страниц
                    if not self.is_large_file or i < 50:
                        page_tables = page.extract_tables()
                        if page_tables:
                            tables.extend(page_tables)
            
            self.metrics["processing_times"]["ocr"] = time.time() - start_time
            self.metrics["file_stats"]["pages"] = pages_count
            self.metrics["file_stats"]["tables_extracted"] = len(tables)
            
            return {
                "text": "\n".join(text_content),
                "tables": tables,
                "metadata": {"pages": pages_count, "large_file_optimized": self.is_large_file},
                "method": "pdfplumber_fallback"
            }
    
    def _process_pdf_text(self) -> Dict[str, Any]:
        """Обрабатывает PDF с текстовым слоем."""
        start_time = time.time()
        
        if DOCLING_AVAILABLE and self.converter:
            # Используем Docling
            result = self.converter.convert(str(self.file_path))
            doc_json = result.document.export_to_dict()
            
            text_content = []
            tables = []
            
            for item in doc_json.get("content", []):
                if item.get("type") == "text":
                    text_content.append(item.get("text", ""))
                elif item.get("type") == "table":
                    tables.append(item)
            
            self.metrics["processing_times"]["text_extraction"] = time.time() - start_time
            self.metrics["file_stats"]["pages"] = len(doc_json.get("pages", []))
            self.metrics["file_stats"]["tables_extracted"] = len(tables)
            
            return {
                "text": "\n".join(text_content),
                "tables": tables,
                "metadata": doc_json.get("metadata", {}),
                "layout": doc_json.get("layout", {}),
                "method": "docling_text"
            }
        else:
            # Fallback: используем pdfplumber
            return self._process_pdf_text_fallback()
    
    def _process_pdf_text_fallback(self) -> Dict[str, Any]:
        """Fallback обработка PDF с текстом."""
        start_time = time.time()
        
        text_content = []
        tables = []
        
        # Для больших файлов обрабатываем постранично с логированием прогресса
        if pdfplumber is None:
            logger.warning("pdfplumber not available, cannot process PDF text")
            return {
                "text": "",
                "tables": [],
                "metadata": {},
                "method": "fallback_no_pdfplumber"
            }
        with pdfplumber.open(str(self.file_path)) as pdf:
            pages_count = len(pdf.pages)
            
            if self.is_large_file:
                logger.info(f"Processing large PDF ({pages_count} pages) - will log progress")
                batch_size = 10  # Обрабатываем по 10 страниц за раз
            
            for i, page in enumerate(pdf.pages):
                # Логирование прогресса для больших файлов
                if self.is_large_file and i % batch_size == 0:
                    logger.info(f"Processing page {i+1}/{pages_count} ({(i+1)/pages_count*100:.1f}%)")
                
                text = page.extract_text()
                if text:
                    text_content.append(f"--- Page {i+1} ---\n{text}")
                
                # Для больших файлов извлекаем таблицы только для первых N страниц или выборочно
                if not self.is_large_file or i < 50:  # Для больших - только первые 50 страниц
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
        
        self.metrics["processing_times"]["text_extraction"] = time.time() - start_time
        self.metrics["file_stats"]["pages"] = pages_count
        self.metrics["file_stats"]["tables_extracted"] = len(tables)
        
        if self.is_large_file:
            logger.info(f"Large PDF processing completed: {pages_count} pages, {len(tables)} tables")
        
        return {
            "text": "\n".join(text_content),
            "tables": tables,
            "metadata": {"pages": pages_count, "large_file_optimized": self.is_large_file},
            "method": "pdfplumber_fallback"
        }
    
    def _process_docx(self) -> Dict[str, Any]:
        """Обрабатывает DOCX файл."""
        start_time = time.time()
        
        if DOCLING_AVAILABLE and self.converter:
            # Используем Docling
            result = self.converter.convert(str(self.file_path))
            doc_json = result.document.export_to_dict()
            
            text_content = []
            tables = []
            
            for item in doc_json.get("content", []):
                if item.get("type") == "text":
                    text_content.append(item.get("text", ""))
                elif item.get("type") == "table":
                    tables.append(item)
            
            self.metrics["processing_times"]["text_extraction"] = time.time() - start_time
            self.metrics["file_stats"]["tables_extracted"] = len(tables)
            
            return {
                "text": "\n".join(text_content),
                "tables": tables,
                "metadata": doc_json.get("metadata", {}),
                "method": "docling_docx"
            }
        else:
            # Fallback: используем python-docx
            return self._process_docx_fallback()
    
    def _process_docx_fallback(self) -> Dict[str, Any]:
        """Fallback обработка DOCX."""
        start_time = time.time()
        
        if Document is None:
            logger.warning("python-docx not available, cannot process DOCX")
            return {
                "text": "",
                "tables": [],
                "metadata": {},
                "method": "fallback_no_docx"
            }
        doc = Document(str(self.file_path))
        
        text_content = []
        tables = []
        
        # Извлечение текста из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Извлечение таблиц
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        self.metrics["processing_times"]["text_extraction"] = time.time() - start_time
        self.metrics["file_stats"]["tables_extracted"] = len(tables)
        
        return {
            "text": "\n".join(text_content),
            "tables": tables,
            "metadata": {},
            "method": "python-docx_fallback"
        }
    
    def _process_html(self) -> Dict[str, Any]:
        """Обрабатывает HTML файл."""
        start_time = time.time()
        
        try:
            if BeautifulSoup is None:
                logger.warning("beautifulsoup4 not available, cannot process HTML")
                return {
                    "text": "",
                    "tables": [],
                    "metadata": {},
                    "method": "fallback_no_bs4"
                }
            with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, "html.parser")
            
            # Извлечение текста
            text_content = soup.get_text(separator="\n", strip=True)
            
            # Извлечение таблиц
            tables = []
            for table in soup.find_all("table"):
                table_data = []
                for row in table.find_all("tr"):
                    row_data = [cell.get_text(strip=True) for cell in row.find_all(["td", "th"])]
                    if row_data:
                        table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            self.metrics["processing_times"]["text_extraction"] = time.time() - start_time
            self.metrics["file_stats"]["tables_extracted"] = len(tables)
            
            return {
                "text": text_content,
                "tables": tables,
                "metadata": {"title": soup.title.string if soup.title else ""},
                "method": "beautifulsoup"
            }
        except Exception as e:
            logger.error(f"Error processing HTML: {e}")
            raise
    
    def _process_fallback(self) -> Dict[str, Any]:
        """Fallback обработка для неизвестных типов."""
        return {
            "text": "",
            "tables": [],
            "metadata": {},
            "method": "fallback"
        }
    
    def _process_with_cloudru_vlm(self) -> Dict[str, Any]:
        """Обработка документа с использованием Cloud.ru VLM (ТОЛЬКО удаленная обработка)"""
        start_time = time.time()
        
        if not self.converter:
            raise Exception("Cloud.ru VLM converter not initialized")
        
        try:
            # Конвертация через удаленную VLM
            logger.info(f"Converting {self.file_info.get('original_name', 'unknown')} with Cloud.ru VLM...")
            result = self.converter.convert(str(self.file_path))
            doc = result.document
            
            # Извлечение данных (ТОЛЬКО из удаленной VLM)
            text_content = doc.export_to_markdown()
            doctags = doc.export_to_document_tokens()
            
            # Извлечение структуры документа
            doc_json = doc.export_to_dict()
            
            # Извлечение таблиц (если есть)
            tables = []
            for item in doc_json.get("content", []):
                if item.get("type") == "table":
                    tables.append(item)
            
            # Извлечение метаданных
            metadata = doc_json.get("metadata", {})
            
            # Извлечение layout информации
            layout = doc_json.get("layout", {})
            
            self.metrics["processing_times"]["vlm"] = time.time() - start_time
            self.metrics["file_stats"]["pages"] = len(doc_json.get("pages", []))
            self.metrics["file_stats"]["tables_extracted"] = len(tables)
            
            return {
                "text": text_content,
                "tables": tables,
                "metadata": metadata,
                "layout": layout,
                "doctags": doctags,
                "method": "cloudru_vlm"
            }
            
        except Exception as e:
            logger.error(f"Error processing with Cloud.ru VLM: {e}")
            raise
    
    def _process_with_cloudru_vlm_s3(self) -> Dict[str, Any]:
        """
        Обработка документа с использованием Cloud.ru VLM через S3 Object Storage.
        PDF → изображения страниц → S3 → pre-signed URLs → Granite-Docling → DocTags/Markdown
        """
        start_time = time.time()
        
        if not self.s3_client:
            raise Exception("Cloud.ru S3 client not initialized")
        
        if not self.cloudru_api_url or not self.cloudru_api_key:
            raise Exception("Cloud.ru API credentials not configured")
        
        try:
            import fitz  # PyMuPDF
            import requests
            import tempfile
            import shutil
            
            logger.info(f"Processing {self.file_info.get('original_name', 'unknown')} with Cloud.ru VLM via S3...")
            
            # 1. Конвертация PDF в изображения страниц
            logger.info("Converting PDF to images...")
            doc = fitz.open(str(self.file_path))
            pages_count = len(doc)
            self.metrics["file_stats"]["pages"] = pages_count
            
            # Создаем временную директорию для изображений
            temp_dir = Path(tempfile.mkdtemp(prefix="docling_s3_"))
            page_urls = []
            
            try:
                for page_num in range(pages_count):
                    page = doc[page_num]
                    
                    # Рендерим страницу в изображение
                    # Уменьшаем размер для лучшей совместимости с Cloud.ru API (1x вместо 2x)
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.0, 1.0))  # 1x для совместимости
                    img_data = pix.tobytes("png")
                    
                    # Сохраняем во временный файл
                    temp_image_path = temp_dir / f"page_{page_num:04d}.png"
                    with open(temp_image_path, "wb") as f:
                        f.write(img_data)
                    
                    # Загружаем в S3
                    object_key = f"{self.unit_id}/pages/page_{page_num:04d}.png"
                    # Используем публичный URL (virtual hosted style)
                    # Pre-signed URLs возвращают 403, публичные URLs доступны (200 OK)
                    presigned_url = self.s3_client.upload_and_get_url(
                        local_path=temp_image_path,
                        object_key=object_key,
                        expires_in=3600,  # 1 час
                        content_type="image/png",
                        use_public_url=True  # Используем публичный URL (virtual hosted style)
                    )
                    
                    if presigned_url:
                        page_urls.append(presigned_url)
                        logger.debug(f"Page {page_num + 1}/{pages_count} uploaded to S3: {object_key}")
                    else:
                        logger.warning(f"Failed to upload page {page_num + 1} to S3")
                
                doc.close()
                
                if not page_urls:
                    raise Exception("No pages were uploaded to S3")
                
                logger.info(f"Uploaded {len(page_urls)} pages to S3, processing with Granite-Docling...")
                
                # 2. Обработка каждой страницы через Cloud.ru Granite-Docling API
                all_text_content = []
                all_tables = []
                
                for page_idx, page_url in enumerate(page_urls):
                    logger.info(f"Processing page {page_idx + 1}/{len(page_urls)}...")
                    
                    # Проверяем доступность URL перед отправкой в API
                    try:
                        url_check = requests.head(page_url, timeout=10, allow_redirects=True)
                        if url_check.status_code != 200:
                            logger.warning(f"URL для страницы {page_idx + 1} недоступен: {url_check.status_code}")
                            all_text_content.append(f"## Page {page_idx + 1}\n\n[Error: URL недоступен ({url_check.status_code})]\n")
                            continue
                        logger.debug(f"URL для страницы {page_idx + 1} доступен (200 OK)")
                    except Exception as e:
                        logger.warning(f"Ошибка проверки URL для страницы {page_idx + 1}: {e}")
                        all_text_content.append(f"## Page {page_idx + 1}\n\n[Error: URL недоступен]\n")
                        continue
                    
                    # Формируем запрос к Cloud.ru API
                    # Используем /v1/completions с URL в тексте промпта (рабочий формат)
                    # vLLM 0.11.0 не поддерживает image_url в /v1/chat/completions, но работает с URL в тексте
                    prompt_text = "Extract all text from this document page. Preserve structure, tables, and formatting. Return in Russian.\n\nImage URL: " + page_url
                    
                    request_data = {
                        "model": self.cloudru_model,
                        "prompt": prompt_text,
                        "max_tokens": 4096,
                        "temperature": 0.0,
                    }
                    
                    # Используем /v1/completions вместо /v1/chat/completions
                    # Заменяем endpoint в URL, если он содержит /v1/chat/completions
                    if "/v1/chat/completions" in self.cloudru_api_url:
                        api_url = self.cloudru_api_url.replace("/v1/chat/completions", "/v1/completions")
                    elif "/v1/completions" not in self.cloudru_api_url:
                        # Если URL не содержит endpoint, добавляем /v1/completions
                        api_url = self.cloudru_api_url.rstrip("/") + "/v1/completions"
                    else:
                        api_url = self.cloudru_api_url
                    
                    # Отправляем запрос с retry логикой
                    max_retries = 3
                    retry_delay = 2
                    response = None
                    
                    for attempt in range(max_retries):
                        try:
                            response = requests.post(
                                api_url,  # Используем /v1/completions
                                json=request_data,
                                headers={
                                    "Authorization": f"Bearer {self.cloudru_api_key}",
                                    "Content-Type": "application/json"
                                },
                                timeout=180
                            )
                            
                            # Если успешно, выходим из цикла retry
                            if response.status_code == 200:
                                break
                            
                            # Если 500 ошибка, пробуем еще раз
                            if response.status_code == 500 and attempt < max_retries - 1:
                                logger.warning(f"500 ошибка для страницы {page_idx + 1}, попытка {attempt + 1}/{max_retries}, повтор через {retry_delay}s...")
                                time.sleep(retry_delay)
                                retry_delay *= 2  # Экспоненциальная задержка
                                continue
                            else:
                                break
                                
                        except requests.exceptions.Timeout:
                            if attempt < max_retries - 1:
                                logger.warning(f"Timeout для страницы {page_idx + 1}, попытка {attempt + 1}/{max_retries}, повтор через {retry_delay}s...")
                                time.sleep(retry_delay)
                                retry_delay *= 2
                            else:
                                logger.error(f"Timeout для страницы {page_idx + 1} после {max_retries} попыток")
                                response = None
                        except Exception as e:
                            logger.error(f"Исключение при запросе для страницы {page_idx + 1}: {e}")
                            response = None
                            break
                    
                    if response.status_code == 200:
                        result = response.json()
                        if 'choices' in result and len(result['choices']) > 0:
                            # /v1/completions возвращает text напрямую, а не message.content
                            page_content = result['choices'][0].get('text', '')
                            all_text_content.append(f"## Page {page_idx + 1}\n\n{page_content}\n")
                            logger.info(f"Page {page_idx + 1} processed successfully ({len(page_content)} chars)")
                        else:
                            logger.warning(f"No content in response for page {page_idx + 1}")
                            logger.debug(f"Response: {result}")
                            all_text_content.append(f"## Page {page_idx + 1}\n\n[Error: No content in response]\n")
                    else:
                        if response:
                            logger.error(f"Error processing page {page_idx + 1}: {response.status_code}")
                            logger.error(f"Response: {response.text[:500]}")
                            all_text_content.append(f"## Page {page_idx + 1}\n\n[Error: {response.status_code}]\n")
                        else:
                            logger.error(f"Error processing page {page_idx + 1}: No response")
                            all_text_content.append(f"## Page {page_idx + 1}\n\n[Error: No response]\n")
                
                # 3. Объединение результатов
                combined_text = "\n".join(all_text_content)
                
                # Извлечение метаданных
                metadata = {
                    "pages": pages_count,
                    "processing_method": "cloudru_vlm_s3",
                    "s3_bucket": self.s3_client.bucket_name,
                    "unit_id": self.unit_id
                }
                
                self.metrics["processing_times"]["vlm_s3"] = time.time() - start_time
                self.metrics["file_stats"]["pages"] = pages_count
                self.metrics["file_stats"]["tables_extracted"] = len(all_tables)
                
                return {
                    "text": combined_text,
                    "tables": all_tables,
                    "metadata": metadata,
                    "layout": {},
                    "doctags": "",  # Можно добавить конвертацию markdown в doctags
                    "method": "cloudru_vlm_s3"
                }
                
            finally:
                # Очистка временных файлов
                if temp_dir.exists():
                    shutil.rmtree(temp_dir, ignore_errors=True)
                    logger.debug(f"Cleaned up temporary directory: {temp_dir}")
            
        except ImportError as e:
            logger.error(f"Required library not available: {e}")
            raise Exception(f"Missing dependency: {e}")
        except Exception as e:
            logger.error(f"Error processing with Cloud.ru VLM via S3: {e}")
            raise
    
    def _save_results(self, result: Dict[str, Any]) -> List[str]:
        """Сохраняет результаты обработки в JSON и Markdown."""
        stem = self.file_path.stem
        json_file = self.output_dir / f"{stem}.json"
        md_file = self.output_dir / f"{stem}.md"
        
        # Сохранение JSON
        json_data = {
            "unit_id": self.unit_id,
            "file": self.file_info.get("original_name", ""),
            "route": self.route,
            "detected_type": self.detected_type,
            "needs_ocr": self.file_info.get("needs_ocr", False),
            "status": "processed",
            "processing_method": result.get("method", "unknown"),
            "text": result.get("text", ""),
            "tables": result.get("tables", []),
            "metadata": result.get("metadata", {}),
            "layout": result.get("layout", {}),
            "metrics": self.metrics
        }
        
        with open(json_file, "w", encoding="utf-8") as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        # Сохранение Markdown
        md_content = f"# {self.file_info.get('original_name', 'Unknown')}\n\n"
        md_content += f"**Unit ID:** {self.unit_id}\n"
        md_content += f"**Route:** {self.route}\n"
        md_content += f"**Type:** {self.detected_type}\n"
        md_content += f"**Needs OCR:** {self.file_info.get('needs_ocr', False)}\n"
        md_content += f"**Method:** {result.get('method', 'unknown')}\n\n"
        
        if result.get("metadata"):
            md_content += "## Metadata\n\n"
            for key, value in result["metadata"].items():
                md_content += f"- **{key}:** {value}\n"
            md_content += "\n"
        
        md_content += "## Content\n\n"
        md_content += result.get("text", "")
        
        if result.get("tables"):
            md_content += "\n\n## Tables\n\n"
            for i, table in enumerate(result["tables"], 1):
                md_content += f"### Table {i}\n\n"
                if isinstance(table, list):
                    # Таблица в виде списка списков
                    for row in table:
                        md_content += "| " + " | ".join(str(cell) for cell in row) + " |\n"
                        if row == table[0]:
                            md_content += "|" + "|".join([" --- "] * len(row)) + "|\n"
                md_content += "\n"
        
        with open(md_file, "w", encoding="utf-8") as f:
            f.write(md_content)
        
        # Обновление размера выходных файлов
        if json_file.exists():
            self.metrics["file_stats"]["output_size"] = json_file.stat().st_size + md_file.stat().st_size
        
        return [str(json_file), str(md_file)]
    
    def _extract_layout(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Извлекает layout структуру документа."""
        layout_start = time.time()
        
        layout = {
            "pages": [],
            "sections": [],
            "blocks": []
        }
        
        # Для PDF - извлекаем информацию о страницах и блоках
        if self.detected_type == "pdf":
            try:
                if pdfplumber is None:
                    return layout
                with pdfplumber.open(str(self.file_path)) as pdf:
                    for i, page in enumerate(pdf.pages):
                        page_layout = {
                            "page_number": i + 1,
                            "width": page.width,
                            "height": page.height,
                            "text_blocks": []
                        }
                        
                        # Извлекаем текстовые блоки через слова
                        words = page.extract_words()
                        if words:
                            page_layout["word_count"] = len(words)
                        
                        layout["pages"].append(page_layout)
            except Exception as e:
                logger.warning(f"Could not extract PDF layout: {e}")
        
        # Для DOCX - извлекаем структуру параграфов и секций
        elif self.detected_type == "docx":
            try:
                if Document is None:
                    return layout
                doc = Document(str(self.file_path))
                sections = []
                for i, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        sections.append({
                            "index": i,
                            "text": paragraph.text[:100],  # Первые 100 символов
                            "style": paragraph.style.name if paragraph.style else "Normal"
                        })
                layout["sections"] = sections
            except Exception as e:
                logger.warning(f"Could not extract DOCX layout: {e}")
        
        # Для HTML - извлекаем структуру тегов
        elif self.detected_type == "html":
            try:
                if BeautifulSoup is None:
                    return layout
                with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, "html.parser")
                
                # Извлекаем основные структурные элементы
                blocks = []
                for tag in soup.find_all(["h1", "h2", "h3", "p", "div"]):
                    if tag.get_text(strip=True):
                        blocks.append({
                            "tag": tag.name,
                            "text_preview": tag.get_text(strip=True)[:100]
                        })
                layout["blocks"] = blocks
            except Exception as e:
                logger.warning(f"Could not extract HTML layout: {e}")
        
        self.metrics["processing_times"]["layout_analysis"] = time.time() - layout_start
        return layout
    
    def _get_stack_trace(self, exception: Exception) -> str:
        """Получает stack trace из исключения."""
        import traceback
        return traceback.format_exc()

